from pathlib import Path
import json
import tempfile


def export_report(report, export_format: str) -> tuple[bytes, str]:
    """Export a report to PDF or DOCX and return (content_bytes, mime_type)."""
    if export_format == "pdf":
        return _export_pdf(report)
    elif export_format == "docx":
        return _export_docx(report)
    raise ValueError(f"Unsupported format: {export_format}")


def _export_pdf(report) -> tuple[bytes, str]:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, report.title, ln=True, align="C")
    pdf.ln(6)

    sections = report.sections if isinstance(report.sections, dict) else {}
    for section_key, content in sections.items():
        # Section heading — strip leading numbering prefix like "01_"
        heading = section_key.split("_", 1)[-1].replace("_", " ").title() if "_" in section_key else section_key

        pdf.set_font("Helvetica", "B", 13)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(0, 9, heading, ln=True, fill=True)
        pdf.ln(2)

        pdf.set_font("Helvetica", "", 10)

        if isinstance(content, str):
            pdf.multi_cell(0, 6, content)
        elif isinstance(content, dict):
            c_type = content.get("type")
            if c_type == "heading":
                pdf.set_font("Helvetica", "B", 11)
                pdf.multi_cell(0, 7, str(content.get("text", "")))
            elif c_type == "analysis":
                pdf.set_font("Helvetica", "I", 10)
                pdf.multi_cell(0, 6, f"Analysis: {content.get('analysis_type', '')}")
                pdf.set_font("Helvetica", "", 9)
                results = content.get("results") or {}
                for k, v in results.items():
                    if not isinstance(v, (dict, list)):
                        val_str = f"{v:.4f}" if isinstance(v, float) else str(v)
                        pdf.multi_cell(0, 5, f"  {k}: {val_str}")
            else:
                pdf.multi_cell(0, 6, json.dumps(content, indent=2)[:500])
        pdf.ln(4)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        pdf.output(str(tmp_path))
        data = tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)

    return data, "application/pdf"


def _export_docx(report) -> tuple[bytes, str]:
    from docx import Document

    doc = Document()
    doc.add_heading(report.title, level=1)

    sections = report.sections if isinstance(report.sections, dict) else {}
    for section_key, content in sections.items():
        heading = section_key.split("_", 1)[-1].replace("_", " ").title() if "_" in section_key else section_key
        doc.add_heading(heading, level=2)

        if isinstance(content, str):
            doc.add_paragraph(content)
        elif isinstance(content, dict):
            c_type = content.get("type")
            if c_type == "heading":
                doc.add_heading(str(content.get("text", "")), level=3)
            elif c_type == "analysis":
                doc.add_paragraph(f"Analysis: {content.get('analysis_type', '')}", style="Intense Quote")
                results = content.get("results") or {}
                for k, v in results.items():
                    if not isinstance(v, (dict, list)):
                        val_str = f"{v:.4f}" if isinstance(v, float) else str(v)
                        doc.add_paragraph(f"{k}: {val_str}", style="List Bullet")
            else:
                doc.add_paragraph(json.dumps(content, indent=2))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        doc.save(str(tmp_path))
        data = tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)

    return data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
